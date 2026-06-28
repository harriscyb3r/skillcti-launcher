"""
File text extraction endpoint.
Accepts PDF and DOCX uploads and returns plain text suitable for passing
to a skill prompt.
"""

import io
from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import JSONResponse

router = APIRouter()

_MAX_BYTES = 20 * 1024 * 1024  # 20 MB hard limit


@router.post("/extract-file")
async def extract_file(file: UploadFile = File(...)):
    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    raw = await file.read()
    if len(raw) > _MAX_BYTES:
        raise HTTPException(413, "File too large (max 20 MB)")

    if ext == "pdf":
        text = _extract_pdf(raw, filename)
    elif ext in ("docx", "doc"):
        text = _extract_docx(raw, filename)
    else:
        # Plain-text types — decode directly
        try:
            text = raw.decode("utf-8", errors="replace")
        except Exception as exc:
            raise HTTPException(422, f"Could not decode file: {exc}") from exc

    if not text.strip():
        raise HTTPException(422, "No text could be extracted from the file")

    return JSONResponse({"text": text, "filename": filename, "chars": len(text)})


def _extract_pdf(raw: bytes, filename: str) -> str:
    try:
        import pdfplumber  # type: ignore
    except ImportError as exc:
        raise HTTPException(500, "pdfplumber not installed — run: pip install pdfplumber") from exc

    try:
        pages = []
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(f"--- Page {i} ---\n{page_text}")
        return "\n\n".join(pages)
    except Exception as exc:
        raise HTTPException(422, f"PDF extraction failed: {exc}") from exc


def _extract_docx(raw: bytes, filename: str) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise HTTPException(500, "python-docx not installed — run: pip install python-docx") from exc

    try:
        doc = Document(io.BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except Exception as exc:
        raise HTTPException(422, f"DOCX extraction failed: {exc}") from exc
